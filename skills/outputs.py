"""
Output generators for skill reports — DOCX and XLSX.
"""

import json
import logging
import os
import time

logger = logging.getLogger(__name__)


def _fmt(val, fallback="N/A"):
    """Format a value for display."""
    if val is None:
        return fallback
    if isinstance(val, float):
        if abs(val) >= 1e9:
            return f"${val / 1e9:.1f}B"
        if abs(val) >= 1e6:
            return f"${val / 1e6:.1f}M"
        if abs(val) < 1:
            return f"{val:.2%}"
        return f"{val:.2f}"
    return str(val)


def _s(val, fallback="N/A"):
    """Safe string conversion."""
    if val is None:
        return fallback
    return str(val)


def _add_executive_snapshot(doc, snapshot):
    """Render the 1-Minute Read executive snapshot at the top of any skill DOCX."""
    if not snapshot or not isinstance(snapshot, dict):
        return
    headline = snapshot.get("headline", "")
    bullets = snapshot.get("bullets", [])
    if not headline and not bullets:
        return

    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    heading = doc.add_heading("1-Minute Read", level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)

    if headline:
        p = doc.add_paragraph()
        run = p.add_run(str(headline))
        run.font.size = Pt(12)
        run.font.bold = True

    for bullet in bullets[:5]:
        p = doc.add_paragraph(str(bullet), style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)

    # Skill-specific verdict/action line
    for key in ("rating_action", "verdict", "score", "recommendation",
                "action_summary", "priorities", "estimated_savings",
                "performance", "key_actions", "outlook"):
        val = snapshot.get(key)
        if val:
            p = doc.add_paragraph()
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x26, 0xA6, 0x9A)
            break

    p = doc.add_paragraph()
    p.add_run("_" * 70).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    p.space_after = Pt(4)


def generate_docx(skill, job, context, output_path):
    """Route to skill-specific DOCX generator."""
    if skill.id == "earnings-analysis":
        _generate_earnings_docx(job, context, output_path)
    else:
        _generate_generic_docx(skill, job, context, output_path)


def generate_xlsx(skill, job, context, output_path):
    """Route to skill-specific XLSX generator."""
    if skill.id == "earnings-analysis":
        _generate_earnings_xlsx(job, context, output_path)
    else:
        _generate_generic_xlsx(skill, job, context, output_path)


# ─── Earnings Analysis DOCX (CL-style report) ────────────────────

def _generate_earnings_docx(job, context, output_path):
    """Generate an institutional-quality earnings update DOCX matching sell-side format."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx not installed, skipping DOCX generation")
        return

    results = context.get("results", {})
    earnings = results.get("earnings", {})
    ratios = results.get("computed_ratios", {})
    company = context.get("data", {}).get("company_info", {})
    header = earnings.get("header", {})
    charts = context.get("charts", [])

    # Build chart lookup by filename prefix
    chart_map = {}
    for c in charts:
        chart_map[c["filename"]] = c

    doc = Document()

    # Page setup — match CL report margins
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Heading styles
    for level in [1, 2, 3]:
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
        hs.font.name = "Calibri"

    ticker = job.inputs.get("ticker", "")
    company_name = header.get("company_name", company.get("name", ticker))
    fig_num = [0]  # mutable counter for figure captions

    def _add_figure(doc, chart_key, caption_override=None):
        """Add a chart image with italic caption."""
        chart = chart_map.get(chart_key)
        if not chart:
            return
        path = chart["path"]
        if not os.path.exists(path):
            return
        fig_num[0] += 1
        doc.add_paragraph("")
        try:
            doc.add_picture(path, width=Inches(5.8))
        except Exception as e:
            logger.warning("Failed to add chart %s: %s", chart_key, e)
            return
        caption = caption_override or chart.get("caption", f"Figure {fig_num[0]}")
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # ─── Title Block ───
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{company_name} ({ticker})")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)

    # Rating line (with bold label and values)
    rating = header.get("rating", earnings.get("rating", "N/A"))
    price = header.get("current_price", "")
    pt = header.get("price_target", "")
    upside = header.get("upside_pct", "")
    p = doc.add_paragraph()
    run = p.add_run("Rating: ")
    run.font.size = Pt(11)
    run = p.add_run(f"{rating}")
    run.font.size = Pt(11)
    run.font.bold = True
    if price:
        run = p.add_run(f"  |  Price: ")
        run.font.size = Pt(11)
        run = p.add_run(f"${price}")
        run.font.size = Pt(11)
        run.font.bold = True
    if pt:
        run = p.add_run(f"  |  Price Target: ")
        run.font.size = Pt(11)
        run = p.add_run(f"${pt}")
        run.font.size = Pt(11)
        run.font.bold = True
    if upside:
        sign = "+" if isinstance(upside, (int, float)) and upside > 0 else ""
        run = p.add_run(f"  |  Upside: {sign}{upside}%")
        run.font.size = Pt(11)

    # Subtitle line
    sector = header.get("sector", company.get("sector", ""))
    industry = header.get("industry", company.get("industry", ""))
    report_date = header.get("report_date", time.strftime("%B %Y"))
    sub = f"Earnings Update  |  {report_date}"
    if sector:
        sub += f"  |  {sector}"
        if industry:
            sub += f" / {industry}"
    p = doc.add_paragraph()
    run = p.add_run(sub)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # ─── Executive Snapshot ───
    _add_executive_snapshot(doc, earnings.get("executive_snapshot", {}))

    # ─── Earnings Summary ───
    es = earnings.get("earnings_summary", {})
    headline = es.get("headline", "Earnings Update")
    doc.add_heading(f"Earnings Summary: {headline}", level=2)

    # Earnings summary table
    est = es.get("table", [])
    if est:
        table = doc.add_table(rows=1 + len(est), cols=5)
        table.style = "Light List Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers_list = ["Metric", "Actual", "Consensus", "Variance", "Result"]
        for i, h in enumerate(headers_list):
            cell = table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
        for j, row_data in enumerate(est):
            row = table.rows[j + 1]
            row.cells[0].text = _s(row_data.get("metric"))
            row.cells[1].text = _s(row_data.get("actual"))
            row.cells[2].text = _s(row_data.get("consensus"))
            row.cells[3].text = _s(row_data.get("variance"))
            row.cells[4].text = _s(row_data.get("result"))
            # Color the result cell
            result_text = _s(row_data.get("result", "")).upper()
            for run in row.cells[4].paragraphs[0].runs:
                if "BEAT" in result_text:
                    run.font.color.rgb = RGBColor(0x26, 0xA6, 0x9A)
                elif "MISS" in result_text:
                    run.font.color.rgb = RGBColor(0xEF, 0x53, 0x50)
        _set_table_font_size(table, Pt(10))

    # GAAP note if present in table
    for row_data in est:
        result_text = _s(row_data.get("result", "")).upper()
        if "*" in result_text or "GAAP" in _s(row_data.get("metric", "")).upper():
            note = row_data.get("variance", "")
            if note and "N/M" in str(note).upper():
                p = doc.add_paragraph()
                run = p.add_run(
                    "*GAAP miss may be driven by one-time items; not indicative of operating performance."
                )
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            break

    # Key Takeaways — with bold lead sentences
    takeaways = es.get("key_takeaways", [])
    if takeaways:
        doc.add_heading("Key Takeaways", level=3)
        for item in takeaways:
            text = str(item)
            p = doc.add_paragraph(style="List Bullet")
            # Try to bold the first sentence (lead sentence)
            colon_pos = text.find(":")
            period_pos = text.find(".")
            split_pos = -1
            if 5 < colon_pos < 80:
                split_pos = colon_pos + 1
            elif 5 < period_pos < 120:
                split_pos = period_pos + 1

            if split_pos > 0:
                lead = text[:split_pos]
                rest = text[split_pos:]
                run = p.add_run(lead)
                run.font.bold = True
                run.font.size = Pt(11)
                if rest.strip():
                    run = p.add_run(f" {rest.strip()}")
                    run.font.size = Pt(11)
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)
        doc.add_paragraph("")

    # Updated Financial Estimates
    ue = earnings.get("updated_estimates", {})
    uet = ue.get("table", [])
    if uet:
        doc.add_heading("Updated Financial Estimates", level=3)
        cols = _get_estimate_columns(uet)
        table = doc.add_table(rows=1 + len(uet), cols=len(cols))
        table.style = "Light List Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(cols):
            cell = table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
        for j, row_data in enumerate(uet):
            row = table.rows[j + 1]
            vals = _get_estimate_row_values(row_data, cols)
            for i, v in enumerate(vals):
                row.cells[i].text = _s(v)
        _set_table_font_size(table, Pt(10))
        source = ue.get("source_note", "")
        if source:
            p = doc.add_paragraph()
            run = p.add_run(source)
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # ─── Detailed Results Analysis ───
    doc.add_heading("Detailed Results Analysis", level=1)

    # Revenue Analysis
    ra = earnings.get("revenue_analysis", {})
    if ra:
        doc.add_heading("Revenue Analysis", level=2)
        if ra.get("narrative"):
            # Split narrative into paragraphs
            for para_text in str(ra["narrative"]).split("\n\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

    # Revenue chart
    _add_figure(doc, "chart1_revenue.png",
                f"Figure {fig_num[0] + 1}: Quarterly Net Sales — Revenue Trend")

    # Segment / Geographic Analysis
    sa = earnings.get("segment_analysis", {})
    if sa:
        title = "Geographic Performance" if _is_geographic(sa) else "Segment Analysis"
        if isinstance(sa, dict):
            doc.add_heading(title, level=2)
            if sa.get("narrative"):
                for para_text in str(sa["narrative"]).split("\n\n"):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
            segs = sa.get("segments", [])
            if segs:
                table = doc.add_table(rows=1 + len(segs), cols=4)
                table.style = "Light List Accent 1"
                for i, h in enumerate(["Segment", "Revenue", "Growth", "Key Drivers"]):
                    table.rows[0].cells[i].text = h
                    for run in table.rows[0].cells[i].paragraphs[0].runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                for j, seg in enumerate(segs):
                    row = table.rows[j + 1]
                    row.cells[0].text = _s(seg.get("name"))
                    row.cells[1].text = _s(seg.get("revenue"))
                    row.cells[2].text = _s(seg.get("growth"))
                    row.cells[3].text = _s(seg.get("key_drivers"))
                _set_table_font_size(table, Pt(10))
        elif isinstance(sa, str):
            doc.add_heading(title, level=2)
            doc.add_paragraph(sa)

    # Segment chart
    _add_figure(doc, "chart3_segments.png",
                f"Figure {fig_num[0] + 1}: Segment / Geographic Growth Breakdown")

    # Profitability Analysis
    pa = earnings.get("profitability_analysis", {})
    if pa:
        doc.add_heading("Profitability Analysis", level=2)
        if pa.get("narrative"):
            for para_text in str(pa["narrative"]).split("\n\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())

    # EPS chart
    _add_figure(doc, "chart2_eps.png",
                f"Figure {fig_num[0] + 1}: EPS — Steady Growth Trend")

    # Gross margin chart
    _add_figure(doc, "chart4_margin.png",
                f"Figure {fig_num[0] + 1}: Gross Margin Trend — Profitability Trajectory")

    # ─── Key Metrics & Guidance ───
    doc.add_heading("Key Metrics & Guidance", level=1)

    # Operating Highlights
    oh = earnings.get("operating_highlights", {})
    if oh:
        doc.add_heading("Operating Highlights", level=2)
        # Source note
        oht = oh.get("table", [])
        if oht:
            cols = _get_dynamic_table_cols(oht, ["metric", "q_value", "fy_value", "prior_fy_value", "yoy_change"],
                                           ["Metric", "Latest Quarter", "FY Current", "FY Prior", "YoY Change"])
            table = doc.add_table(rows=1 + len(oht), cols=len(cols))
            table.style = "Light List Accent 1"
            for i, (key, label) in enumerate(cols):
                table.rows[0].cells[i].text = label
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for j, row_data in enumerate(oht):
                row = table.rows[j + 1]
                for i, (key, _) in enumerate(cols):
                    row.cells[i].text = _s(row_data.get(key))
            _set_table_font_size(table, Pt(10))

        # Source note for operating highlights
        p = doc.add_paragraph()
        run = p.add_run(f"Source: {company_name} Earnings Release.")
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        if oh.get("free_cash_flow_narrative"):
            doc.add_paragraph(oh["free_cash_flow_narrative"])
        if oh.get("capital_allocation"):
            doc.add_paragraph(oh["capital_allocation"])

    # FCF chart
    _add_figure(doc, "chart5_fcf.png",
                f"Figure {fig_num[0] + 1}: Key Operating Metrics — Year-over-Year")

    # Guidance
    gu = earnings.get("guidance", {})
    if gu:
        doc.add_heading("Forward Guidance", level=2)
        if gu.get("narrative"):
            doc.add_paragraph(gu["narrative"])
        gut = gu.get("table", [])
        if gut:
            # Determine columns — check for scenario_framework/notes
            has_notes = any(r.get("notes") or r.get("scenario_framework") for r in gut)
            col_headers = ["Metric", "Guidance Range"]
            if has_notes:
                col_headers.append("Notes")
            table = doc.add_table(rows=1 + len(gut), cols=len(col_headers))
            table.style = "Light List Accent 1"
            for i, h in enumerate(col_headers):
                table.rows[0].cells[i].text = h
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for j, row_data in enumerate(gut):
                row = table.rows[j + 1]
                row.cells[0].text = _s(row_data.get("metric"))
                row.cells[1].text = _s(row_data.get("guidance_range"))
                if has_notes:
                    notes = row_data.get("notes") or row_data.get("scenario_framework", "")
                    row.cells[2].text = _s(notes)
            _set_table_font_size(table, Pt(10))

        # Source note
        p = doc.add_paragraph()
        run = p.add_run(f"Source: {company_name} Guidance (Earnings Call)")
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # Estimates chart
    _add_figure(doc, "chart7_estimates.png",
                f"Figure {fig_num[0] + 1}: Revenue — Forward Estimates")

    # ─── Updated Investment Thesis ───
    it = earnings.get("investment_thesis", {})
    if it:
        doc.add_heading("Updated Investment Thesis", level=1)

        # Thesis Pillars
        pillars = it.get("pillars", [])
        if pillars:
            doc.add_heading("Thesis Pillar Assessment", level=2)
            for pillar in pillars:
                name = pillar.get("name", "")
                status = pillar.get("status", "")
                detail = pillar.get("detail", "")
                p = doc.add_paragraph()
                run = p.add_run(f"{name}: ")
                run.font.bold = True
                run.font.size = Pt(11)
                run2 = p.add_run(status)
                run2.font.bold = True
                run2.font.size = Pt(11)
                if "STRENGTHEN" in status.upper():
                    run2.font.color.rgb = RGBColor(0x26, 0xA6, 0x9A)
                elif "WEAKEN" in status.upper():
                    run2.font.color.rgb = RGBColor(0xEF, 0x53, 0x50)
                if detail:
                    p2 = doc.add_paragraph()
                    run = p2.add_run(detail)
                    run.font.size = Pt(10)

        # Key Risks
        risks = it.get("key_risks", [])
        if risks:
            doc.add_heading("Key Risks", level=2)
            for risk in risks:
                text = str(risk)
                p = doc.add_paragraph(style="List Bullet")
                # Bold lead (up to colon)
                colon_pos = text.find(":")
                if 3 < colon_pos < 60:
                    run = p.add_run(text[:colon_pos + 1])
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run = p.add_run(text[colon_pos + 1:])
                    run.font.size = Pt(11)
                else:
                    run = p.add_run(text)
                    run.font.size = Pt(11)

    # ─── Valuation & Estimates ───
    val = earnings.get("valuation", {})
    if val:
        doc.add_heading("Valuation & Estimates", level=1)

        # Peer Comparison Table
        peers = val.get("peer_table", [])
        if peers:
            doc.add_heading("Peer Valuation Comparison", level=2)
            if val.get("narrative"):
                doc.add_paragraph(val["narrative"])
            # Source note
            p = doc.add_paragraph()
            run = p.add_run("Source: Yahoo Finance, StockAnalysis.com. *Estimated from most recent reported quarter.")
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

            table = doc.add_table(rows=1 + len(peers), cols=6)
            table.style = "Light List Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(["Company", "Market Cap", "Forward P/E", "EV/EBITDA", "Growth", "Gross Margin"]):
                table.rows[0].cells[i].text = h
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for j, peer in enumerate(peers):
                row = table.rows[j + 1]
                row.cells[0].text = _s(peer.get("company"))
                row.cells[1].text = _s(peer.get("market_cap"))
                row.cells[2].text = _s(peer.get("forward_pe"))
                row.cells[3].text = _s(peer.get("ev_ebitda"))
                row.cells[4].text = _s(peer.get("growth"))
                row.cells[5].text = _s(peer.get("gross_margin"))
            _set_table_font_size(table, Pt(10))

        # Peer P/E chart
        _add_figure(doc, "chart6_peer_pe.png",
                    f"Figure {fig_num[0] + 1}: Forward P/E — Peer Valuation Comparison")

        # Price Target Methodology
        ptm = val.get("price_target_methodology", [])
        if ptm:
            doc.add_heading("Price Target Methodology", level=2)
            pt_text = header.get("price_target", val.get("blended_target", ""))
            if pt_text:
                doc.add_paragraph(f"Our ${pt_text} price target is derived from a blended approach:")

            table = doc.add_table(rows=1 + len(ptm) + 1, cols=4)
            table.style = "Light List Accent 1"
            for i, h in enumerate(["Methodology", "Implied Value", "Weight", "Contribution"]):
                table.rows[0].cells[i].text = h
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for j, m in enumerate(ptm):
                row = table.rows[j + 1]
                row.cells[0].text = _s(m.get("methodology"))
                iv = m.get("implied_value")
                row.cells[1].text = f"${iv}" if iv and "$" not in str(iv) else _s(iv)
                w = m.get("weight")
                row.cells[2].text = f"{w}%" if w and "%" not in str(w) else _s(w)
                row.cells[3].text = f"${m.get('contribution')}" if m.get("contribution") and "$" not in str(m.get("contribution")) else _s(m.get("contribution"))
            # Blended target row
            last_row = table.rows[-1]
            last_row.cells[0].text = "Blended Price Target"
            bt = val.get("blended_target", "N/A")
            last_row.cells[3].text = f"${bt}" if bt and "$" not in str(bt) else _s(bt)
            for cell in last_row.cells:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
            _set_table_font_size(table, Pt(10))

        # Detailed Estimate Updates
        ue2 = earnings.get("updated_estimates", {})
        uet2 = ue2.get("table", [])
        if uet2:
            doc.add_heading("Detailed Estimate Updates", level=2)
            cols = _get_estimate_columns(uet2)
            table = doc.add_table(rows=1 + len(uet2), cols=len(cols))
            table.style = "Light List Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(cols):
                cell = table.rows[0].cells[i]
                cell.text = h
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for j, row_data in enumerate(uet2):
                row = table.rows[j + 1]
                vals = _get_estimate_row_values(row_data, cols)
                for i, v in enumerate(vals):
                    row.cells[i].text = _s(v)
            _set_table_font_size(table, Pt(10))
            p = doc.add_paragraph()
            src = ue2.get("source_note", f"Source: {company_name} filings; estimates from analyst consensus. A = Actual, E = Estimate.")
            run = p.add_run(src)
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        # PT methodology chart
        _add_figure(doc, "chart8_pt.png",
                    f"Figure {fig_num[0] + 1}: Price Target Methodology — Blended Approach")

    # ─── Appendix ───
    appendix = earnings.get("appendix", {})
    if appendix:
        doc.add_heading("Appendix", level=1)

        highlights = appendix.get("earnings_call_highlights", [])
        if highlights:
            doc.add_heading("Earnings Call Highlights", level=2)
            p = doc.add_paragraph()
            run = p.add_run("Key quotes from the earnings call:")
            run.font.bold = True
            run.font.size = Pt(11)
            for h in highlights:
                p = doc.add_paragraph()
                run = p.add_run(str(h))
                run.font.size = Pt(10)
                run.font.italic = True

        notable = appendix.get("notable_items")
        if notable:
            doc.add_heading("Notable Items", level=2)
            doc.add_paragraph(str(notable))

        analyst = appendix.get("analyst_activity")
        if analyst:
            doc.add_heading("Post-Earnings Analyst Activity", level=2)
            p = doc.add_paragraph()
            run = p.add_run(str(analyst))
            run.font.size = Pt(9)
            run.font.italic = True

    # ─── Sources & References ───
    doc.add_heading("Sources & References", level=1)
    sources = [
        f"{company_name} Quarterly Earnings Release",
        "Earnings Call Transcript",
        f"{company_name} Investor Relations",
        f"{ticker} Analyst Forecasts — StockAnalysis.com",
        f"Yahoo Finance — {ticker}",
    ]
    for src in sources:
        p = doc.add_paragraph()
        run = p.add_run(f"  {src}")
        run.font.size = Pt(10)

    # ─── Disclaimer ───
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(
        "DISCLAIMER: This report is for informational purposes only and does not constitute "
        "investment advice. All estimates are based on publicly available information and AI analysis "
        "as of the date indicated. Past performance is not indicative of future results. "
        "Investors should conduct their own research before making investment decisions."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.save(output_path)
    logger.info("Generated earnings DOCX: %s", output_path)


def _is_geographic(sa):
    """Detect if segment analysis is geographic rather than product-based."""
    if not isinstance(sa, dict):
        return False
    segments = sa.get("segments", [])
    geo_terms = {"america", "europe", "asia", "africa", "latin", "north", "pacific",
                 "emea", "apac", "international", "domestic", "region"}
    for seg in segments:
        name = str(seg.get("name", "")).lower()
        if any(term in name for term in geo_terms):
            return True
    return False


def _get_estimate_columns(table_data):
    """Determine column headers for estimates table from data."""
    if not table_data:
        return ["Metric"]
    first = table_data[0]
    cols = ["Metric"]
    # Check for labeled columns
    if first.get("fy_actual_label"):
        cols.append(first["fy_actual_label"])
    elif "fy_actual" in first:
        cols.append("FY Actual")
    if first.get("fy1_label"):
        cols.append(first["fy1_label"])
    elif "fy1_estimate" in first:
        cols.append("FY+1 Est")
    if first.get("fy2_label"):
        cols.append(first["fy2_label"])
    elif "fy2_estimate" in first:
        cols.append("FY+2 Est")
    return cols


def _get_estimate_row_values(row_data, cols):
    """Extract values for an estimates table row."""
    vals = [row_data.get("metric", "")]
    key_map = {1: "fy_actual", 2: "fy1_estimate", 3: "fy2_estimate"}
    for i in range(1, len(cols)):
        vals.append(row_data.get(key_map.get(i, ""), ""))
    return vals


def _get_dynamic_table_cols(table_data, keys, labels):
    """Build column list from data, only including columns that have data."""
    if not table_data:
        return list(zip(keys[:1], labels[:1]))
    cols = []
    for key, label in zip(keys, labels):
        if any(row.get(key) is not None for row in table_data):
            cols.append((key, label))
    return cols if cols else list(zip(keys[:1], labels[:1]))


def _set_table_font_size(table, size):
    """Set font size for all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = size


# ─── Earnings Analysis XLSX ───────────────────────────────────

def _generate_earnings_xlsx(job, context, output_path):
    """Generate a detailed earnings analysis Excel workbook."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        logger.error("openpyxl not installed, skipping XLSX generation")
        return

    results = context.get("results", {})
    earnings = results.get("earnings", {})
    ratios = results.get("computed_ratios", {})
    header = earnings.get("header", {})
    company = context.get("data", {}).get("company_info", {})

    wb = Workbook()
    hdr_font = Font(bold=True, size=14, color="17365D")
    sub_font = Font(bold=True, size=11, color="17365D")
    tbl_hdr_fill = PatternFill(start_color="17365D", end_color="17365D", fill_type="solid")
    tbl_hdr_font = Font(bold=True, size=10, color="FFFFFF")
    normal_font = Font(size=10)
    beat_font = Font(size=10, color="26A69A", bold=True)
    miss_font = Font(size=10, color="EF5350", bold=True)
    thin_border = Border(
        bottom=Side(style="thin", color="D1D4DC"),
    )

    ticker = job.inputs.get("ticker", "")
    company_name = header.get("company_name", company.get("name", ticker))

    # ─── Sheet 1: Summary ───
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_properties.tabColor = "17365D"

    ws.merge_cells("A1:E1")
    ws["A1"] = f"{company_name} ({ticker}) — Earnings Update"
    ws["A1"].font = hdr_font

    rating = header.get("rating", earnings.get("rating", ""))
    price = header.get("current_price", "")
    pt_val = header.get("price_target", "")
    ws["A2"] = f"Rating: {rating}  |  Price: ${price}  |  PT: ${pt_val}"
    ws["A2"].font = Font(size=10, color="777777")

    row = 4
    ws.cell(row=row, column=1, value="Earnings Summary").font = sub_font
    row += 1

    # Earnings summary table
    es = earnings.get("earnings_summary", {})
    est = es.get("table", [])
    if est:
        for i, h in enumerate(["Metric", "Actual", "Consensus", "Variance", "Result"], 1):
            c = ws.cell(row=row, column=i, value=h)
            c.font = tbl_hdr_font
            c.fill = tbl_hdr_fill
        row += 1
        for rd in est:
            ws.cell(row=row, column=1, value=rd.get("metric", "")).font = normal_font
            ws.cell(row=row, column=2, value=rd.get("actual", "")).font = normal_font
            ws.cell(row=row, column=3, value=rd.get("consensus", "")).font = normal_font
            ws.cell(row=row, column=4, value=rd.get("variance", "")).font = normal_font
            result = rd.get("result", "")
            c = ws.cell(row=row, column=5, value=result)
            c.font = beat_font if "BEAT" in str(result).upper() else (miss_font if "MISS" in str(result).upper() else normal_font)
            row += 1
        row += 1

    # Key ratios
    if ratios:
        ws.cell(row=row, column=1, value="Key Financial Ratios").font = sub_font
        row += 1
        for i, h in enumerate(["Metric", "Value"], 1):
            c = ws.cell(row=row, column=i, value=h)
            c.font = tbl_hdr_font
            c.fill = tbl_hdr_fill
        row += 1
        for k, v in ratios.items():
            if v is not None:
                ws.cell(row=row, column=1, value=k.replace("_", " ").title()).font = normal_font
                ws.cell(row=row, column=2, value=v).font = normal_font
                row += 1
        row += 1

    # Updated estimates
    ue = earnings.get("updated_estimates", {})
    uet = ue.get("table", [])
    if uet:
        ws.cell(row=row, column=1, value="Updated Estimates").font = sub_font
        row += 1
        cols = _get_estimate_columns(uet)
        for i, h in enumerate(cols, 1):
            c = ws.cell(row=row, column=i, value=h)
            c.font = tbl_hdr_font
            c.fill = tbl_hdr_fill
        row += 1
        for rd in uet:
            vals = _get_estimate_row_values(rd, cols)
            for i, v in enumerate(vals, 1):
                ws.cell(row=row, column=i, value=v).font = normal_font
            row += 1
        row += 1

    # Peer valuation
    val = earnings.get("valuation", {})
    peers = val.get("peer_table", [])
    if peers:
        ws.cell(row=row, column=1, value="Peer Valuation").font = sub_font
        row += 1
        peer_headers = ["Company", "Market Cap", "Fwd P/E", "EV/EBITDA", "Growth", "Gross Margin"]
        for i, h in enumerate(peer_headers, 1):
            c = ws.cell(row=row, column=i, value=h)
            c.font = tbl_hdr_font
            c.fill = tbl_hdr_fill
        row += 1
        for peer in peers:
            ws.cell(row=row, column=1, value=peer.get("company", "")).font = normal_font
            ws.cell(row=row, column=2, value=peer.get("market_cap", "")).font = normal_font
            ws.cell(row=row, column=3, value=peer.get("forward_pe", "")).font = normal_font
            ws.cell(row=row, column=4, value=peer.get("ev_ebitda", "")).font = normal_font
            ws.cell(row=row, column=5, value=peer.get("growth", "")).font = normal_font
            ws.cell(row=row, column=6, value=peer.get("gross_margin", "")).font = normal_font
            row += 1
        row += 1

    # Operating highlights
    oh = earnings.get("operating_highlights", {})
    oht = oh.get("table", [])
    if oht:
        ws.cell(row=row, column=1, value="Operating Highlights").font = sub_font
        row += 1
        dyn_cols = _get_dynamic_table_cols(oht,
            ["metric", "q_value", "fy_value", "prior_fy_value", "yoy_change"],
            ["Metric", "Latest Qtr", "FY Current", "FY Prior", "YoY Change"])
        for i, (_, label) in enumerate(dyn_cols, 1):
            c = ws.cell(row=row, column=i, value=label)
            c.font = tbl_hdr_font
            c.fill = tbl_hdr_fill
        row += 1
        for rd in oht:
            for i, (key, _) in enumerate(dyn_cols, 1):
                ws.cell(row=row, column=i, value=rd.get(key, "")).font = normal_font
            row += 1
        row += 1

    # Guidance
    gu = earnings.get("guidance", {})
    gut = gu.get("table", [])
    if gut:
        ws.cell(row=row, column=1, value="Forward Guidance").font = sub_font
        row += 1
        for i, h in enumerate(["Metric", "Guidance Range", "Notes"], 1):
            c = ws.cell(row=row, column=i, value=h)
            c.font = tbl_hdr_font
            c.fill = tbl_hdr_fill
        row += 1
        for rd in gut:
            ws.cell(row=row, column=1, value=rd.get("metric", "")).font = normal_font
            ws.cell(row=row, column=2, value=rd.get("guidance_range", "")).font = normal_font
            ws.cell(row=row, column=3, value=rd.get("notes", "")).font = normal_font
            row += 1

    # Auto-width
    _auto_width(ws)

    # ─── Sheet 2: Income Statements ───
    inc_annual = context.get("data", {}).get("income_statements", {})
    inc_quarterly = context.get("data", {}).get("income_statements_quarterly", {})
    if inc_annual or inc_quarterly:
        ws2 = wb.create_sheet("Income Statements")
        ws2.sheet_properties.tabColor = "26A69A"
        _write_financial_sheet(ws2, "Annual Income Statement", inc_annual, 1, tbl_hdr_font, tbl_hdr_fill, normal_font)
        start = len(inc_annual) + 5 if inc_annual else 1
        _write_financial_sheet(ws2, "Quarterly Income Statement", inc_quarterly, start, tbl_hdr_font, tbl_hdr_fill, normal_font)
        _auto_width(ws2)

    # ─── Sheet 3: Balance Sheet ───
    bs = context.get("data", {}).get("balance_sheet", {})
    if bs:
        ws3 = wb.create_sheet("Balance Sheet")
        ws3.sheet_properties.tabColor = "FF9800"
        _write_financial_sheet(ws3, "Balance Sheet", bs, 1, tbl_hdr_font, tbl_hdr_fill, normal_font)
        _auto_width(ws3)

    # ─── Sheet 4: Cash Flow ───
    cf = context.get("data", {}).get("cash_flow", {})
    if cf:
        ws4 = wb.create_sheet("Cash Flow")
        ws4.sheet_properties.tabColor = "AB47BC"
        _write_financial_sheet(ws4, "Cash Flow Statement", cf, 1, tbl_hdr_font, tbl_hdr_fill, normal_font)
        _auto_width(ws4)

    # ─── Sheet 5: Price Data ───
    prices = context.get("data", {}).get("prices", [])
    if prices:
        ws5 = wb.create_sheet("Price Data")
        ws5.sheet_properties.tabColor = "2962FF"
        ws5.append(["Date", "Open", "High", "Low", "Close", "Volume"])
        for i in range(1, 7):
            ws5.cell(row=1, column=i).font = tbl_hdr_font
            ws5.cell(row=1, column=i).fill = tbl_hdr_fill
        for p in prices:
            ws5.append([p["date"], p["open"], p["high"], p["low"], p["close"], p["volume"]])
        _auto_width(ws5)

    wb.save(output_path)
    logger.info("Generated earnings XLSX: %s", output_path)


def _write_financial_sheet(ws, title, data, start_row, hdr_font, hdr_fill, normal_font):
    """Write a financial statement (income/balance/cash flow) data block to a sheet."""
    if not data:
        return
    ws.cell(row=start_row, column=1, value=title).font = Font(bold=True, size=12, color="17365D")
    start_row += 1

    periods = sorted(data.keys(), reverse=True)
    # Header row
    ws.cell(row=start_row, column=1, value="Line Item").font = hdr_font
    ws.cell(row=start_row, column=1).fill = hdr_fill
    for i, period in enumerate(periods, 2):
        c = ws.cell(row=start_row, column=i, value=period)
        c.font = hdr_font
        c.fill = hdr_fill

    # Get all line items from first period
    if periods:
        line_items = list(data[periods[0]].keys())
        for j, item in enumerate(line_items):
            row = start_row + 1 + j
            ws.cell(row=row, column=1, value=item).font = normal_font
            for i, period in enumerate(periods, 2):
                val = data.get(period, {}).get(item)
                if val is not None:
                    ws.cell(row=row, column=i, value=val).font = normal_font


def _auto_width(ws):
    """Auto-size column widths."""
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 3, 55)


# ─── Generic DOCX (for non-earnings skills) ──────────────────

def _generate_generic_docx(skill, job, context, output_path):
    """Generate a generic Word document report for non-earnings skills."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed, skipping DOCX generation")
        return

    doc = Document()

    title = doc.add_heading(f"{skill.name}: {job.inputs.get('ticker', '')}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Generated: {context['results'].get('_meta', {}).get('completed_at', 'N/A')} | "
        f"Skill: {skill.id}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    results = context.get("results", {})

    # ─── Executive Snapshot ───
    snapshot = None
    for rkey in results:
        if isinstance(results[rkey], dict) and "executive_snapshot" in results[rkey]:
            snapshot = results[rkey]["executive_snapshot"]
            break
    _add_executive_snapshot(doc, snapshot)

    # Company Info
    company = context.get("data", {}).get("company_info", {})
    if company:
        doc.add_heading("Company Overview", level=1)
        table = doc.add_table(rows=4, cols=2)
        table.style = "Light List Accent 1"
        _set_cell(table, 0, 0, "Company")
        _set_cell(table, 0, 1, company.get("name", "N/A"))
        _set_cell(table, 1, 0, "Sector / Industry")
        _set_cell(table, 1, 1, f"{company.get('sector', 'N/A')} / {company.get('industry', 'N/A')}")
        _set_cell(table, 2, 0, "Market Cap")
        _set_cell(table, 2, 1, _fmt(company.get("market_cap")))
        _set_cell(table, 3, 0, "Country")
        _set_cell(table, 3, 1, company.get("country", "N/A"))
        doc.add_paragraph("")

    # Financial Health section
    health = results.get("financial_health", {})
    if health and "raw_response" not in health:
        doc.add_heading("Financial Health Assessment", level=1)
        doc.add_paragraph(
            f"Overall Score: {health.get('overall_score', 'N/A')}/10 — "
            f"{health.get('overall_rating', 'N/A')}"
        )
        for section in ["liquidity", "solvency", "profitability", "efficiency", "growth"]:
            sub = health.get(section, {})
            if isinstance(sub, dict) and sub.get("score"):
                doc.add_heading(section.title(), level=2)
                for k, v in sub.items():
                    if k != "score":
                        doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}", style="List Bullet")
                doc.add_paragraph(f"Score: {sub.get('score', 'N/A')}/10")
        _add_list_section(doc, "Red Flags", health.get("red_flags", []))
        _add_list_section(doc, "Strengths", health.get("strengths", []))
        if health.get("summary"):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(health["summary"])

    # DCF section
    dcf = results.get("dcf", {})
    if dcf and "raw_response" not in dcf:
        doc.add_heading("DCF Valuation", level=1)
        doc.add_paragraph(f"WACC: {_fmt(dcf.get('wacc_estimate'))} | "
                         f"Terminal Growth: {_fmt(dcf.get('terminal_growth_rate'))}")
        doc.add_paragraph(f"Implied Share Price: {_fmt(dcf.get('implied_share_price'))} | "
                         f"Current: {_fmt(dcf.get('current_price'))}")
        doc.add_paragraph(f"Upside/Downside: {_fmt(dcf.get('upside_downside_pct'))}%")
        doc.add_paragraph(f"Verdict: {dcf.get('valuation_verdict', 'N/A')}")
        if dcf.get("summary"):
            doc.add_paragraph(dcf["summary"])

    # Comps section
    comps = results.get("comps", {})
    if comps and "raw_response" not in comps:
        doc.add_heading("Comparable Analysis", level=1)
        doc.add_paragraph(f"Average Implied Price: {_fmt(comps.get('average_implied_price'))}")
        doc.add_paragraph(f"Verdict: {comps.get('relative_verdict', 'N/A')}")
        if comps.get("summary"):
            doc.add_paragraph(comps["summary"])

    # Earnings section (fallback for old-format data)
    earnings = results.get("earnings", {})
    if earnings and "raw_response" not in earnings:
        doc.add_heading("Earnings Analysis", level=1)
        doc.add_paragraph(f"Score: {earnings.get('overall_score', 'N/A')}/10 — "
                         f"{earnings.get('rating', 'N/A')}")
        if earnings.get("summary"):
            doc.add_paragraph(earnings["summary"])

    # Sector section
    sector = results.get("sector", {})
    if sector and "raw_response" not in sector:
        doc.add_heading("Sector Analysis", level=1)
        doc.add_paragraph(f"Sector Score: {sector.get('overall_sector_score', 'N/A')}/10")
        if sector.get("summary"):
            doc.add_paragraph(sector["summary"])

    # QC section
    qc = results.get("quality_check", {})
    if qc:
        doc.add_heading("Quality Check", level=2)
        doc.add_paragraph(f"Passed: {'Yes' if qc.get('passed') else 'No'} | "
                         f"Confidence: {qc.get('confidence_score', 'N/A')}/10")
        _add_list_section(doc, "Issues", qc.get("issues", []))

    # Computed ratios
    ratios = results.get("computed_ratios", {})
    if ratios:
        doc.add_heading("Key Financial Ratios", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        _set_cell(table, 0, 0, "Metric")
        _set_cell(table, 0, 1, "Value")
        for k, v in ratios.items():
            if v is not None:
                row = table.add_row()
                row.cells[0].text = k.replace("_", " ").title()
                row.cells[1].text = _fmt(v)

    doc.save(output_path)
    logger.info("Generated DOCX: %s", output_path)


# ─── Generic XLSX (for non-earnings skills) ──────────────────

def _generate_generic_xlsx(skill, job, context, output_path):
    """Generate a generic Excel workbook for non-earnings skills."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
    except ImportError:
        logger.error("openpyxl not installed, skipping XLSX generation")
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"

    header_font = Font(bold=True, size=14)
    subheader_font = Font(bold=True, size=11)

    ws["A1"] = f"{skill.name}: {job.inputs.get('ticker', '')}"
    ws["A1"].font = header_font
    ws["A2"] = f"Generated: {context['results'].get('_meta', {}).get('completed_at', '')}"

    results = context.get("results", {})
    row = 4

    ratios = results.get("computed_ratios", {})
    if ratios:
        ws.cell(row=row, column=1, value="Key Metrics").font = subheader_font
        row += 1
        for k, v in ratios.items():
            if v is not None:
                ws.cell(row=row, column=1, value=k.replace("_", " ").title())
                ws.cell(row=row, column=2, value=v if isinstance(v, (int, float)) else str(v))
                row += 1
        row += 1

    health = results.get("financial_health", {})
    if health and "raw_response" not in health:
        ws.cell(row=row, column=1, value="Financial Health").font = subheader_font
        row += 1
        ws.cell(row=row, column=1, value="Overall Score")
        ws.cell(row=row, column=2, value=health.get("overall_score", ""))
        row += 1
        ws.cell(row=row, column=1, value="Rating")
        ws.cell(row=row, column=2, value=health.get("overall_rating", ""))
        row += 1
        for section in ["liquidity", "solvency", "profitability", "efficiency", "growth"]:
            sub = health.get(section, {})
            if isinstance(sub, dict) and sub.get("score"):
                ws.cell(row=row, column=1, value=f"  {section.title()} Score")
                ws.cell(row=row, column=2, value=sub["score"])
                row += 1
        row += 1

    dcf = results.get("dcf", {})
    if dcf and "raw_response" not in dcf:
        ws.cell(row=row, column=1, value="DCF Valuation").font = subheader_font
        row += 1
        for key in ["wacc_estimate", "terminal_growth_rate", "enterprise_value",
                     "equity_value", "implied_share_price", "current_price", "upside_downside_pct"]:
            val = dcf.get(key)
            if val is not None:
                ws.cell(row=row, column=1, value=key.replace("_", " ").title())
                ws.cell(row=row, column=2, value=val)
                row += 1
        row += 1

    prices = context.get("data", {}).get("prices", [])
    if prices:
        ws2 = wb.create_sheet("Price Data")
        ws2.append(["Date", "Open", "High", "Low", "Close", "Volume"])
        for p in prices:
            ws2.append([p["date"], p["open"], p["high"], p["low"], p["close"], p["volume"]])

    for sheet in wb.sheetnames:
        _auto_width(wb[sheet])

    wb.save(output_path)
    logger.info("Generated XLSX: %s", output_path)


def _set_cell(table, row, col, text):
    """Set text in a docx table cell."""
    table.rows[row].cells[col].text = str(text)


def _add_list_section(doc, title, items):
    """Add a bulleted list section to a docx document."""
    if not items:
        return
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")
